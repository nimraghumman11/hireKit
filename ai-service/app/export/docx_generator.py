"""Generate a polished DOCX from interview kit data using python-docx."""
from __future__ import annotations
import logging
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from config import get_settings

logger = logging.getLogger(__name__)

# Brand colours
BRAND_HEX   = "4361EE"   # indigo-600
MUTED_HEX   = "64748B"   # slate-500
DARK_HEX    = "0F172A"   # slate-900
SUCCESS_HEX = "10B981"   # emerald-500
WARN_HEX    = "F59E0B"   # amber-500
ERROR_HEX   = "EF4444"   # red-500

EASY_HEX   = "10B981"
MEDIUM_HEX = "F59E0B"
HARD_HEX   = "EF4444"


def _hex(h: str) -> RGBColor:
    r, g, b = int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)
    return RGBColor(r, g, b)


def _set_cell_bg(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _safe(value, fallback: str = "") -> str:
    if value is None:
        return fallback
    return str(value)


def _add_heading(doc: Document, text: str, level: int = 1, color_hex: str = BRAND_HEX) -> None:
    para = doc.add_heading(text, level=level)
    para.runs[0].font.color.rgb = _hex(color_hex)


def _add_labeled_para(doc: Document, label: str, text: str) -> None:
    para = doc.add_paragraph()
    run_label = para.add_run(f"{label}: ")
    run_label.bold = True
    run_label.font.size = Pt(10)
    run_label.font.color.rgb = _hex(MUTED_HEX)
    run_text = para.add_run(text)
    run_text.font.size = Pt(10)


def _add_bullet_list(doc: Document, items: list, indent_level: int = 0) -> None:
    for item in items:
        para = doc.add_paragraph(style="List Bullet")
        para.paragraph_format.left_indent = Inches(0.25 * (indent_level + 1))
        run = para.add_run(_safe(item))
        run.font.size = Pt(10)


def _add_separator(doc: Document) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.space_before = Pt(0)
    run = para.add_run("─" * 80)
    run.font.size = Pt(6)
    run.font.color.rgb = _hex("E2E8F0")


def generate_docx(kit_data: dict, kit_id: str) -> str:
    settings = get_settings()
    output_dir = Path(settings.pdf_storage_path) / "kits" / kit_id
    output_dir.mkdir(parents=True, exist_ok=True)
    docx_path = output_dir / "interview-kit.docx"

    role_title     = _safe(kit_data.get("roleTitle", "Role"))
    department     = _safe(kit_data.get("department", ""))
    exp_level      = _safe(kit_data.get("experienceLevel", ""))
    jd             = kit_data.get("jobDescription") or {}
    behavioral     = kit_data.get("behavioralQuestions") or []
    technical      = kit_data.get("technicalQuestions") or []
    scorecard      = kit_data.get("scorecard") or []
    rubric         = kit_data.get("rubric") or []
    lang_issues    = kit_data.get("languageIssues") or []

    doc = Document()

    # ── Page margins ──────────────────��──────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Cover / Title ────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(role_title)
    title_run.bold = True
    title_run.font.size = Pt(26)
    title_run.font.color.rgb = _hex(BRAND_HEX)

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run(f"{department}  ·  {exp_level}")
    sub_run.font.size = Pt(12)
    sub_run.font.color.rgb = _hex(MUTED_HEX)

    doc.add_paragraph()
    _add_separator(doc)
    doc.add_paragraph()

    # ── Job Description ──────────────────────────────────────
    if jd:
        _add_heading(doc, "Job Description", level=1)

        if jd.get("summary"):
            _add_heading(doc, "Summary", level=2, color_hex=MUTED_HEX)
            p = doc.add_paragraph(_safe(jd["summary"]))
            p.runs[0].font.size = Pt(10)

        if jd.get("aboutTheRole"):
            _add_heading(doc, "About the Role", level=2, color_hex=MUTED_HEX)
            p = doc.add_paragraph(_safe(jd["aboutTheRole"]))
            p.runs[0].font.size = Pt(10)

        if jd.get("whyJoinUs"):
            _add_heading(doc, "Why Join Us", level=2, color_hex=MUTED_HEX)
            p = doc.add_paragraph(_safe(jd["whyJoinUs"]))
            p.runs[0].font.size = Pt(10)

        if jd.get("responsibilities"):
            _add_heading(doc, "What You'll Do", level=2, color_hex=MUTED_HEX)
            _add_bullet_list(doc, jd["responsibilities"])

        if jd.get("requiredQualifications"):
            _add_heading(doc, "Required Qualifications", level=2, color_hex=MUTED_HEX)
            _add_bullet_list(doc, jd["requiredQualifications"])
        elif jd.get("requiredSkills"):
            _add_heading(doc, "Required Skills", level=2, color_hex=MUTED_HEX)
            _add_bullet_list(doc, jd["requiredSkills"])

        if jd.get("preferredQualifications"):
            _add_heading(doc, "Preferred Qualifications", level=2, color_hex=MUTED_HEX)
            _add_bullet_list(doc, jd["preferredQualifications"])
        elif jd.get("niceToHaveSkills"):
            _add_heading(doc, "Nice-to-Have Skills", level=2, color_hex=MUTED_HEX)
            _add_bullet_list(doc, jd["niceToHaveSkills"])

        if jd.get("whatYoullBring"):
            _add_heading(doc, "What You'll Bring", level=2, color_hex=MUTED_HEX)
            p = doc.add_paragraph(_safe(jd["whatYoullBring"]))
            p.runs[0].font.size = Pt(10)

        if jd.get("compensationAndBenefits"):
            _add_heading(doc, "Compensation & Benefits", level=2, color_hex=MUTED_HEX)
            p = doc.add_paragraph(_safe(jd["compensationAndBenefits"]))
            p.runs[0].font.size = Pt(10)

        doc.add_page_break()

    # ── Behavioral Questions ─────────────────────────────────
    if behavioral:
        _add_heading(doc, f"Behavioral Questions  ({len(behavioral)})", level=1)

        for i, q in enumerate(behavioral, 1):
            para = doc.add_paragraph()
            run_num = para.add_run(f"Q{i}. ")
            run_num.bold = True
            run_num.font.color.rgb = _hex(BRAND_HEX)
            run_q = para.add_run(_safe(q.get("question")))
            run_q.bold = True
            run_q.font.size = Pt(10)

            if q.get("competency"):
                _add_labeled_para(doc, "Competency", q["competency"])

            if q.get("evalCriteria"):
                _add_labeled_para(doc, "Evaluation Criteria", q["evalCriteria"])

            if q.get("followUps"):
                p = doc.add_paragraph()
                p.add_run("Follow-up Questions:").bold = True
                _add_bullet_list(doc, q["followUps"], indent_level=1)

            sg = q.get("scoringGuide") or {}
            if sg:
                p = doc.add_paragraph()
                p.add_run("Scoring Guide:").bold = True
                score_labels = {"1": "Poor (1)", "3": "Meets Expectations (3)", "5": "Exceeds Expectations (5)"}
                for score_key in ("1", "3", "5"):
                    if sg.get(score_key):
                        inner = doc.add_paragraph(style="List Bullet")
                        inner.paragraph_format.left_indent = Inches(0.5)
                        run_k = inner.add_run(f"{score_labels.get(score_key, score_key)}: ")
                        run_k.bold = True
                        run_k.font.size = Pt(9)
                        inner.add_run(_safe(sg[score_key])).font.size = Pt(9)

            doc.add_paragraph()

        doc.add_page_break()

    # ── Technical Questions ──────────────────────────────────
    if technical:
        _add_heading(doc, f"Technical Questions  ({len(technical)})", level=1)
        diff_colors = {"easy": EASY_HEX, "medium": MEDIUM_HEX, "hard": HARD_HEX}

        for i, q in enumerate(technical, 1):
            para = doc.add_paragraph()
            run_num = para.add_run(f"Q{i}. ")
            run_num.bold = True
            run_num.font.color.rgb = _hex(BRAND_HEX)
            run_q = para.add_run(_safe(q.get("question")))
            run_q.bold = True
            run_q.font.size = Pt(10)

            diff = _safe(q.get("difficulty", "medium")).lower()
            diff_color = diff_colors.get(diff, MUTED_HEX)
            meta = doc.add_paragraph()
            run_d = meta.add_run(f"[{diff.upper()}]  ")
            run_d.bold = True
            run_d.font.color.rgb = _hex(diff_color)
            run_d.font.size = Pt(9)
            if q.get("topic"):
                run_t = meta.add_run(f"Topic: {q['topic']}")
                run_t.font.size = Pt(9)
                run_t.font.color.rgb = _hex(MUTED_HEX)

            if q.get("evalCriteria"):
                _add_labeled_para(doc, "Evaluation Criteria", q["evalCriteria"])

            if q.get("sampleAnswer"):
                _add_labeled_para(doc, "Sample Answer", q["sampleAnswer"])

            if q.get("redFlags"):
                p = doc.add_paragraph()
                run_rf = p.add_run("Red Flags: ")
                run_rf.bold = True
                run_rf.font.color.rgb = _hex(ERROR_HEX)
                _add_bullet_list(doc, q["redFlags"], indent_level=1)

            doc.add_paragraph()

        doc.add_page_break()

    # ── Scorecard ────────────────────────────────────────────
    if scorecard:
        _add_heading(doc, "Interview Scorecard", level=1)

        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for cell, text in zip(hdr, ["Competency", "Weight", "Score (1-5)", "Notes"]):
            cell.text = text
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            _set_cell_bg(cell, BRAND_HEX.replace("#", ""))

        for row in scorecard:
            cells = table.add_row().cells
            cells[0].text = _safe(row.get("competency"))
            cells[1].text = f"{float(row.get('weight', 0)):.0%}"
            cells[2].text = _safe(row.get("score"), "—")
            cells[3].text = _safe(row.get("notes", ""))
            for cell in cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)

        doc.add_paragraph()
        doc.add_page_break()

    # ── Skills Rubric ────────────────────────────────────────
    if rubric:
        _add_heading(doc, "Skills Rubric", level=1)

        for item in rubric:
            _add_heading(doc, _safe(item.get("skill")), level=2, color_hex=DARK_HEX)
            levels = item.get("proficiencyLevels") or {}
            level_colors = {
                "novice": MUTED_HEX,
                "intermediate": "3B82F6",
                "advanced": BRAND_HEX,
                "expert": "7C3AED",
            }
            for level, desc in levels.items():
                p = doc.add_paragraph()
                run_l = p.add_run(f"{level.capitalize()}: ")
                run_l.bold = True
                run_l.font.color.rgb = _hex(level_colors.get(level, MUTED_HEX))
                run_l.font.size = Pt(9)
                run_d = p.add_run(_safe(desc))
                run_d.font.size = Pt(9)
            doc.add_paragraph()

    # ── Language & Inclusion Review ──────────────────────────
    _add_heading(doc, "Language & Inclusion Review", level=1)
    if lang_issues:
        source_labels = {
            "jobDescription": "Job Description",
            "behavioralQuestions": "Behavioral Questions",
            "technicalQuestions": "Technical Questions",
            "rubric": "Skills Rubric",
        }
        sev_colors = {"error": ERROR_HEX, "warning": WARN_HEX}
        for issue in sorted(lang_issues, key=lambda x: 0 if x.get("severity") == "error" else 1):
            p = doc.add_paragraph()
            sev = _safe(issue.get("severity", "warning")).lower()
            label = f"[{sev.upper()}] "
            run_sev = p.add_run(label)
            run_sev.bold = True
            run_sev.font.color.rgb = _hex(sev_colors.get(sev, MUTED_HEX))
            run_sev.font.size = Pt(9)
            run_text = p.add_run(
                f'"{issue.get("term")}"  →  {issue.get("suggestion")}'
            )
            run_text.font.size = Pt(9)
            source = source_labels.get(issue.get("source", ""), issue.get("source", ""))
            src_p = doc.add_paragraph(f"Found in: {source}")
            src_p.runs[0].font.size = Pt(8)
            src_p.runs[0].font.color.rgb = _hex(MUTED_HEX)
            src_p.paragraph_format.space_after = Pt(4)
    else:
        p = doc.add_paragraph("No inclusive language issues detected.")
        p.runs[0].font.color.rgb = _hex(SUCCESS_HEX)

    doc.save(str(docx_path))
    logger.info("DOCX saved to %s", docx_path)

    base_url = get_settings().ai_service_base_url.rstrip("/")
    return f"{base_url}/static/pdfs/kits/{kit_id}/interview-kit.docx"
